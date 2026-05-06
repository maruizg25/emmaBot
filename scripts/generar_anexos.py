# scripts/generar_anexos.py
# Genera docs/ANEXOS_IT-DIO-2026-206_SercoBot.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores SERCOP ────────────────────────────────────────────
AZUL_SERCOP = RGBColor(0x00, 0x33, 0x66)
AZUL_CLARO  = RGBColor(0x1F, 0x6E, 0xB0)
GRIS_TABLA  = RGBColor(0xD9, 0xE1, 0xF2)
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO       = RGBColor(0x00, 0x00, 0x00)
VERDE       = RGBColor(0x00, 0x60, 0x00)
ROJO        = RGBColor(0xC0, 0x00, 0x00)
NARANJA     = RGBColor(0xBF, 0x60, 0x00)

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color="BBBBBB", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"),  sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = AZUL_SERCOP
    return p


def subheading(text):
    return heading(text, level=2)


def body(text, size=10, bold=False, italic=False, color=NEGRO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def code_line(text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def add_table(headers, rows, col_widths=None, alt_rows=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, AZUL_SERCOP)
        set_cell_border(c, color="FFFFFF")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        tr = t.add_row()
        bg = RGBColor(0xEE, 0xF4, 0xFF) if (alt_rows and ri % 2 == 0) else BLANCO
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            set_cell_border(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(val, tuple):
                text, color = val
                run = p.add_run(text)
                run.font.color.rgb = color
            else:
                run = p.add_run(str(val))
                run.font.color.rgb = NEGRO
            run.font.size = Pt(8.5)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in t.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    return t


def annexe_banner(num, title):
    """Barra azul oscuro de encabezado de anexo."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.rows[0].cells[0]
    set_cell_bg(c, AZUL_SERCOP)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  ANEXO {num}  |  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLANCO
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════

t_cover = doc.add_table(rows=1, cols=1)
t_cover.style = "Table Grid"
c = t_cover.rows[0].cells[0]
set_cell_bg(c, AZUL_SERCOP)
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "SERCOP — Servicio Nacional de Contratacion Publica\n"
    "Coordinacion de Tecnologia de la Informacion\n\n"
    "ANEXOS\n"
    "Informe Tecnico IT-DIO-2026-206\n"
    "SercoBot — Asistente Virtual de Contratacion Publica\n\n"
    "Fecha: 16 de abril de 2026\n"
    "Version: 1.0"
)
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = BLANCO

doc.add_paragraph()

add_table(
    ["Anexo", "Titulo", "Pagina"],
    [
        ["Anexo 1", "Arquitectura tecnica detallada del sistema SercoBot",              "2"],
        ["Anexo 2", "Listado completo de documentos normativos indexados",              "4"],
        ["Anexo 3", "Informe de diagnostico de conectividad (13/04/2026)",              "5"],
        ["Anexo 4", "Ticket MANTIS — Habilitacion de infraestructura (15/04/2026)",    "7"],
        ["Anexo 5", "Resultados de pruebas funcionales — casos de uso validados",      "9"],
    ],
    col_widths=[2.0, 12.0, 2.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANEXO 1 — ARQUITECTURA TECNICA DETALLADA
# ═══════════════════════════════════════════════════════════════

annexe_banner(1, "Arquitectura tecnica detallada del sistema SercoBot")

subheading("1.1 Diagrama de componentes")

body(
    "El sistema SercoBot esta compuesto por dos capas independientes que se integran "
    "a traves del modulo brain.py. La capa de mensajeria (WhatsApp) permanece estable "
    "y desacoplada de la capa de conocimiento normativo (RAG)."
)

code_line("┌─────────────────────────────────────────────────────────────┐")
code_line("│            CAPA DE MENSAJERIA (WhatsApp)                    │")
code_line("│                                                             │")
code_line("│  Meta Cloud API                                             │")
code_line("│    └─> POST /webhook  [agent/main.py]                       │")
code_line("│          ├─> parsear_webhook()  [providers/meta.py]         │")
code_line("│          ├─> deduplicacion (_mensajes_procesados set)       │")
code_line("│          ├─> [multimedia?] ──> _responder_multimedia()      │")
code_line("│          └─> background: _procesar_mensaje()                │")
code_line("│                ├─> obtener_historial()  [memory.py]         │")
code_line("│                ├─> generar_respuesta()  [brain.py]  ───┐    │")
code_line("│                ├─> guardar_mensaje()    [memory.py]    │    │")
code_line("│                └─> enviar_mensaje()     [meta.py]      │    │")
code_line("│                                                        │    │")
code_line("├────────────────────────────────────────────────────────┼────┤")
code_line("│            CAPA DE CONOCIMIENTO (RAG + Tools)          │    │")
code_line("│                                                         ▼    │")
code_line("│  brain.py::generar_respuesta()                              │")
code_line("│    ├─> [shortcut?]  ──> respuesta fija (0ms, 0 tokens)      │")
code_line("│    ├─> LLM 1ra llamada: ?tool_calls                         │")
code_line("│    │     ├─> obtener_montos_pie()     [tools.py]            │")
code_line("│    │     ├─> recomendar_tipo_contratacion() [tools.py]      │")
code_line("│    │     ├─> obtener_plazos(tipo)     [tools.py]            │")
code_line("│    │     ├─> info_rup()               [tools.py]            │")
code_line("│    │     └─> obtener_fecha_hora_ecuador() [tools.py]        │")
code_line("│    ├─> buscar_contexto(query)         [retriever.py]        │")
code_line("│    │     ├─> expand_query (siglas SERCOP)                   │")
code_line("│    │     ├─> embed (nomic-embed-text 768d via Ollama)        │")
code_line("│    │     ├─> HNSW coseno pgvector (top-12 semantico)        │")
code_line("│    │     ├─> GIN tsvector español (top-N lexico)            │")
code_line("│    │     ├─> RRF fusion (top-12 unicos)                     │")
code_line("│    │     └─> cross-encoder reranker (top-4 final)           │")
code_line("│    └─> LLM 2da llamada: gemma4:e2b (Ollama)                 │")
code_line("│          └─> respuesta con citacion de articulo             │")
code_line("│                                                             │")
code_line("├─────────────────────────────────────────────────────────────┤")
code_line("│            INFRAESTRUCTURA                                  │")
code_line("│                                                             │")
code_line("│  PostgreSQL 16 + pgvector 0.8.2  (192.168.2.2:5432)        │")
code_line("│    ├─ tabla: documentos (17 registros)                      │")
code_line("│    ├─ tabla: chunks (3,059 registros + embeddings 768d)     │")
code_line("│    └─ tabla: mensajes (historial de conversaciones)         │")
code_line("│                                                             │")
code_line("│  Ollama  (localhost:11434)                                  │")
code_line("│    ├─ gemma4:e2b  — LLM principal (~5s/respuesta)           │")
code_line("│    └─ nomic-embed-text — embeddings 768d                    │")
code_line("│                                                             │")
code_line("│  FastAPI + Uvicorn  (192.168.9.230:8000)                    │")
code_line("│  Citrix reverse proxy  (192.168.100.131)                    │")
code_line("│  NAT perimetral  (157.100.62.125:443)                       │")
code_line("│  DNS: sercobot.sercop.gob.ec                                │")
code_line("└─────────────────────────────────────────────────────────────┘")

doc.add_paragraph()
subheading("1.2 Flujo de red — publicacion institucional")

add_table(
    ["Capa", "Componente", "IP / Host", "Puerto"],
    [
        ["DNS publico",         "sercobot.sercop.gob.ec",           "157.100.62.125",    "443 HTTPS"],
        ["Firewall / NAT",      "Firewall perimetral SERCOP",        "157.100.62.125",    "443 → interno"],
        ["DMZ red 100",         "IP interna asignada por TI",        "192.168.100.131",   "443"],
        ["Citrix reverse proxy","Publicacion institucional",          "192.168.100.131",   "443 → 8000"],
        ["Servidor aplicacion", "FastAPI SercoBot",                   "192.168.9.230",     "8000"],
        ["Base de datos RAG",   "PostgreSQL 16 + pgvector",          "192.168.2.2",       "5432"],
        ["Motor IA",            "Ollama (gemma4:e2b + nomic-embed)", "localhost",         "11434"],
    ],
    col_widths=[3.5, 5.0, 4.5, 3.5]
)

doc.add_paragraph()
subheading("1.3 Stack tecnologico")

add_table(
    ["Componente", "Tecnologia", "Version", "Rol"],
    [
        ["Servidor API",        "FastAPI + Uvicorn",              "0.104+ / 0.24+",  "Webhook WhatsApp y endpoints admin"],
        ["LLM produccion",      "gemma4:e2b via Ollama",          "Gemma 4 / v0.20+","Generacion de respuestas en lenguaje natural"],
        ["Embeddings",          "nomic-embed-text via Ollama",    "768 dimensiones", "Vectorizacion de chunks y queries"],
        ["Vector DB",           "PostgreSQL 16 + pgvector",       "16 / 0.8.2",      "Busqueda hibrida semantica + lexica"],
        ["Reranker",            "cross-encoder mmarco-mMiniLMv2", "L12-H384-v1",     "Re-ordenamiento de chunks por relevancia"],
        ["Mensajeria",          "Meta Cloud API",                 "v21.0",           "Recepcion y envio de mensajes WhatsApp"],
        ["BD conversaciones",   "PostgreSQL 16 (asyncpg)",        "16",              "Historial de conversaciones por usuario"],
        ["Sistema operativo",   "RHEL 9.3",                       "9.3",             "Servidor de produccion"],
        ["Lenguaje",            "Python",                         "3.11",            "Backend completo"],
    ],
    col_widths=[3.5, 5.0, 3.0, 5.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANEXO 2 — DOCUMENTOS NORMATIVOS INDEXADOS
# ═══════════════════════════════════════════════════════════════

annexe_banner(2, "Listado completo de documentos normativos indexados")

body(
    "La base de conocimiento de SercoBot contiene 3,059 chunks distribuidos en 17 documentos. "
    "El 91.5% de los chunks incluye metadata de articulo para citacion legal precisa."
)

doc.add_paragraph()

add_table(
    ["#", "Documento", "Tipo", "Chunks", "% con articulo"],
    [
        ["1",  "Reglamento General LOSNCP — octubre 2025",              "reglamento",  "1,132", "100%"],
        ["2",  "Normativa Secundaria de Contratacion Publica 2025",     "reglamento",  "1,125", "100%"],
        ["3",  "LOSNCP — RO 140 07-X-2025",                            "ley",         "113",   "100%"],
        ["4",  "Modelo de Pliego SICAE",                                "resolucion",  "141",   "97%"],
        ["5",  "Metodologia de Control Final Sumillada",                "resolucion",  "137",   "97%"],
        ["6",  "Codigo de Etica institucional",                         "resolucion",  "89",    "97%"],
        ["7",  "Manual SOCE — Subasta Inversa Electronica",             "manual_soce", "108",   "1%"],
        ["8",  "RE-SERCOP-2026-0001 (Metodologia de Control)",         "resolucion",  "60",    "97%"],
        ["9",  "RE-SERCOP-2026-0002 (Comite Interinstitucional)",      "resolucion",  "42",    "97%"],
        ["10", "Norma Interna SICAE (farmacos y bienes estrategicos)",  "resolucion",  "33",    "97%"],
        ["11", "Instructivo Extorsion",                                 "resolucion",  "29",    "97%"],
        ["12", "Resolucion Regimen de Transicion — nov. 2025",         "resolucion",  "14",    "97%"],
        ["13", "Manual SOCE — Fase Contractual Bienes y Servicios",    "manual_soce", "19",    "1%"],
        ["14", "RE-SERCOP-2024-0144",                                  "resolucion",  "5",     "97%"],
        ["15", "RE-SERCOP-2025-0152",                                  "resolucion",  "5",     "97%"],
        ["16", "Fe de Errata RE-SERCOP-2024-0142",                    "resolucion",  "1",     "97%"],
        ["17", "Glosario de Terminos SERCOP",                          "manual",      "6",     "0%"],
        ["",   "TOTAL",                                                 "17 documentos","3,059","91.5%"],
    ],
    col_widths=[0.8, 8.5, 2.5, 1.8, 2.9]
)

doc.add_paragraph()
body("Nota: Los manuales SOCE (manual_soce) tienen 1% de metadata de articulo porque son documentos procedimentales, no normativos. Es el comportamiento esperado.", italic=True)

doc.add_paragraph()
subheading("Documentos pendientes de indexar")

add_table(
    ["Documento", "Tipo", "Prioridad", "Por que importa"],
    [
        ["Manual SOCE — Menor Cuantia bienes/servicios", "manual_soce", "Alta",   "Proceso mas comun en contratacion publica"],
        ["Manual SOCE — Menor Cuantia obras",           "manual_soce", "Alta",   "Segundo proceso mas frecuente"],
        ["Manual SOCE — Registro de Proveedores",       "manual_soce", "Alta",   "Alta demanda de ciudadanos nuevos"],
        ["Manual SOCE — Contrataciones de Emergencia",  "manual_soce", "Media",  "Preguntas frecuentes en situaciones urgentes"],
        ["Manual SOCE — Feria Inclusiva",               "manual_soce", "Media",  "Audiencia EPS y MIPYMES"],
        ["COA (Codigo Organico Administrativo)",        "ley",         "Media",  "Capitulos de recursos e impugnaciones formales"],
        ["Resolucion montos PIE 2026",                  "resolucion",  "Alta",   "Fuente citable oficial para umbrales USD 2026"],
    ],
    col_widths=[5.5, 2.5, 2.0, 6.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANEXO 3 — INFORME DE DIAGNOSTICO DE CONECTIVIDAD
# ═══════════════════════════════════════════════════════════════

annexe_banner(3, "Informe de diagnostico de conectividad (13/04/2026)")

body("Script ejecutado: scripts/detectar_firewall.py  |  Equipo: JRUI-L9790-UIO  |  Fecha: 2026-04-13 09:46")
doc.add_paragraph()

subheading("Servicios locales")
add_table(
    ["Servicio", "Host:Puerto", "Estado", "Observacion"],
    [
        ["Ollama (LLM + embeddings)",    "localhost:11434", ("BLOQUEADO", ROJO),   "Servicio no activo en ese momento — ahora en produccion"],
        ["PostgreSQL (base RAG)",        "localhost:5432",  ("BLOQUEADO", ROJO),   "Servicio no activo en ese momento — ahora en VM pg-db (192.168.2.2)"],
        ["FastAPI (servidor del bot)",   "localhost:8000",  ("BLOQUEADO", ROJO),   "Servicio no activo en ese momento — ahora corriendo en 192.168.9.230:8000"],
    ],
    col_widths=[4.5, 3.0, 2.5, 6.5]
)

doc.add_paragraph()
subheading("Meta WhatsApp API")
add_table(
    ["Servicio", "Host:Puerto", "Estado", "Observacion"],
    [
        ["graph.facebook.com HTTPS", "graph.facebook.com:443", ("PARCIAL / CRITICO", NARANJA), "SSL rechazado por Citrix DPI — pendiente bypass Franklin Arias"],
        ["graph.facebook.com DNS",   "graph.facebook.com:443", ("ABIERTO", VERDE),             "Resolucion DNS correcta (28ms)"],
    ],
    col_widths=[4.5, 3.5, 3.0, 5.5]
)

doc.add_paragraph()
subheading("Portal SERCOP y acceso a internet")
add_table(
    ["Servicio", "Host:Puerto", "Estado", "Latencia"],
    [
        ["compraspublicas.gob.ec",       "portal.compraspublicas.gob.ec:443", ("ABIERTO", VERDE),    "27ms — HTTP 200"],
        ["sercop.gob.ec",                "www.sercop.gob.ec:443",            ("ABIERTO", VERDE),    "16ms — HTTP 200"],
        ["PyPI (dependencias Python)",   "pypi.org:443",                     ("ABIERTO", VERDE),    "41ms — HTTP 200"],
        ["Docker Hub",                   "registry-1.docker.io:443",         ("ABIERTO", VERDE),    "110ms"],
        ["Ollama Download (modelos LLM)","registry.ollama.ai:443",           ("ABIERTO", VERDE),    "22ms"],
        ["DNS publico Google",           "8.8.8.8:53",                       ("BLOQUEADO", ROJO),   "DROP silencioso"],
        ["HuggingFace (reranker)",       "huggingface.co:443",               ("PARCIAL", NARANJA),  "SSL rechazado — bypass habilitado 15/04"],
    ],
    col_widths=[4.5, 4.5, 2.5, 5.0]
)

doc.add_paragraph()
subheading("Resumen y acciones resultantes")
add_table(
    ["Metrica", "Valor"],
    [
        ["Total pruebas realizadas",    "12"],
        ["Conexiones abiertas",         "6 (50%)"],
        ["Conexiones bloqueadas",       "4 (33%)"],
        ["Conexiones parciales",        "2 (17%)"],
        ["Criticos sin resolver al 13/04", "4 servicios (Ollama local, PostgreSQL local, Meta HTTPS, DNS externo)"],
        ["Estado al 15/04/2026",        "Resuelto: Ollama + PostgreSQL en produccion. Pendiente: graph.facebook.com SSL bypass"],
    ],
    col_widths=[6.5, 10.0]
)

doc.add_paragraph()
body(
    "Accion pendiente critica: solicitar a Franklin Arias (Administrador de Firewall) "
    "agregar graph.facebook.com a la lista de bypass de inspeccion SSL del firewall perimetral. "
    "Sin esta habilitacion el bot genera respuestas pero no puede entregarlas al usuario via Meta API.",
    bold=True, color=ROJO
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANEXO 4 — TICKET MANTIS HABILITACION INFRAESTRUCTURA
# ═══════════════════════════════════════════════════════════════

annexe_banner(4, "Ticket MANTIS — Habilitacion de infraestructura (15/04/2026)")

body("Sistema: MANTIS Bug Tracker — SERCOP  |  Area: Infraestructura y Operaciones TIC")
doc.add_paragraph()

subheading("Datos del ticket")
add_table(
    ["Campo", "Valor"],
    [
        ["Numero de ticket",   "Pendiente asignacion MANTIS"],
        ["Titulo",             "Habilitacion de infraestructura para SercoBot — Asistente Virtual SERCOP"],
        ["Proyecto",           "TIC — Infraestructura y Operaciones"],
        ["Categoria",          "Solicitud de servicio"],
        ["Prioridad",          "Alta"],
        ["Estado final",       "Resuelto (parcialmente — ver pendientes)"],
        ["Fecha apertura",     "2026-04-13"],
        ["Fecha resolucion",   "2026-04-15"],
        ["Solicitante",        "Mauricio Ruiz — Analista de Desarrollo TIC"],
        ["Asignado a",         "Rolando Coello (DNS) / Franklin Arias (Firewall) / Hugo Yepez (Autorizacion)"],
    ],
    col_widths=[5.0, 11.5]
)

doc.add_paragraph()
subheading("Hilo de atencion — cronologia")

add_table(
    ["Fecha / Hora", "Responsable", "Accion realizada", "Estado"],
    [
        ["2026-04-13 09:46", "Mauricio Ruiz",  "Ejecucion script diagnostico_firewall.py — identificacion de bloqueos criticos",                    ("Completado", VERDE)],
        ["2026-04-13 10:00", "Mauricio Ruiz",  "Envio de solicitud formal al equipo TIC con reporte de firewall adjunto",                           ("Completado", VERDE)],
        ["2026-04-14",       "Equipo TIC",     "Revision de solicitud — aprobacion por Hugo Yepez",                                                  ("Completado", VERDE)],
        ["2026-04-15 17:34", "Hugo Yepez",     "Autorizacion formal del proyecto SercoBot para publicacion institucional",                           ("Completado", VERDE)],
        ["2026-04-15 19:48", "Franklin Arias", "Habilitacion reglas NAT perimetral: puertos 443 y 80 desde 157.100.62.125 hacia 192.168.100.131",   ("Completado", VERDE)],
        ["2026-04-15 20:16", "Rolando Coello", "Configuracion DNS interno: sercobot.sercop.gob.ec -> 192.168.100.131",                              ("Completado", VERDE)],
        ["2026-04-15 20:30", "Mauricio Ruiz",  "Verificacion publica: curl https://sercobot.sercop.gob.ec/ -> HTTP 200",                           ("Completado", VERDE)],
        ["Pendiente",        "Franklin Arias", "Agregar graph.facebook.com a bypass de inspeccion SSL (Citrix DPI)",                                ("Pendiente", NARANJA)],
        ["Pendiente",        "Franklin Arias", "Agregar github.com a bypass SSL para habilitar git pull en servidor",                               ("Pendiente", NARANJA)],
    ],
    col_widths=[2.5, 3.5, 8.5, 2.0]
)

doc.add_paragraph()
subheading("Reglas de firewall habilitadas el 15/04/2026 (Franklin Arias)")

add_table(
    ["#", "Servicio", "Host destino", "Puerto", "Uso"],
    [
        ["1",  "Groq API",           "api.groq.com",                  "443", "LLM en la nube (respaldo)"],
        ["2",  "Meta WhatsApp",      "graph.facebook.com",            "443", "Envio y recepcion de mensajes WhatsApp Business"],
        ["3",  "Cloudflare Tunnel",  "*.cloudflare.com",              "443", "Tunel alternativo de publicacion"],
        ["4",  "GitHub",             "github.com",                    "443", "Control de versiones del codigo"],
        ["5",  "GitHub objetos",     "objects.githubusercontent.com", "443", "Descarga de releases y assets de GitHub"],
        ["6",  "PyPI",               "pypi.org",                      "443", "Instalacion de dependencias Python"],
        ["7",  "PyPI archivos",      "files.pythonhosted.org",        "443", "Descarga de paquetes Python"],
        ["8",  "Ollama",             "ollama.com / registry.ollama.ai","443", "Descarga de modelos LLM"],
        ["9",  "HuggingFace",        "huggingface.co",                "443", "Descarga del modelo reranker cross-encoder"],
        ["10", "HuggingFace CDN",    "cdn-lfs.huggingface.co",        "443", "CDN de archivos de modelos HuggingFace"],
    ],
    col_widths=[0.7, 3.5, 5.3, 1.5, 5.5]
)

doc.add_paragraph()
subheading("Infraestructura provisionada")

add_table(
    ["Componente", "Detalle", "Responsable", "Fecha"],
    [
        ["Dominio publico",          "sercobot.sercop.gob.ec",                              "Rolando Coello", "2026-04-15"],
        ["IP publica",               "157.100.62.125 — NAT perimetral puertos 443 y 80",   "Franklin Arias", "2026-04-15"],
        ["IP interna DMZ (red 100)", "192.168.100.131 / mask 255.255.255.0 / GW .35",      "Rolando Coello", "2026-04-15"],
        ["Servidor aplicacion",      "192.168.9.230 — puerto 8000 (FastAPI)",               "Mauricio Ruiz",  "2026-04-15"],
        ["Publicacion",              "Citrix reverse proxy HTTP->HTTPS",                    "Franklin Arias", "2026-04-15"],
        ["Base de datos RAG",        "PostgreSQL 16 + pgvector 0.8.2 (192.168.2.2)",       "Mauricio Ruiz",  "2026-04-14"],
        ["Autorizacion institucional","Hugo Yepez — 2026-04-15 17:34",                     "Hugo Yepez",     "2026-04-15"],
    ],
    col_widths=[4.0, 6.5, 3.5, 2.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANEXO 5 — RESULTADOS DE PRUEBAS FUNCIONALES
# ═══════════════════════════════════════════════════════════════

annexe_banner(5, "Resultados de pruebas funcionales — casos de uso validados")

body(
    "Las pruebas funcionales se realizaron el 15 y 16 de abril de 2026 sobre el entorno de "
    "produccion (sercobot.sercop.gob.ec). Se validaron los 9 escenarios criticos definidos "
    "en el documento de prueba de escritorio PED-DIO-2026-206."
)

doc.add_paragraph()

subheading("5.1 Tabla de resultados por caso de uso")

add_table(
    ["N.", "Caso de uso", "Entrada de prueba", "Resultado esperado", "Resultado obtenido", "Estado"],
    [
        ["CU-01",
         "Saludo puro -> menu bienvenida",
         '"hola"',
         "Menu de bienvenida con 5 categorias. Sin llamada al LLM.",
         "Recibido menu completo en < 500ms. Sin consumo de tokens.",
         ("PASA", VERDE)],

        ["CU-02",
         "Pregunta normativa -> RAG + articulo",
         '"causales de suspension del RUP"',
         "Respuesta con causales y citacion Art. 16-18 LOSNCP.",
         "Respuesta correcta con citacion Art. 16 LOSNCP. RAG devolvio 4 chunks relevantes.",
         ("PASA", VERDE)],

        ["CU-03",
         "Monto infima cuantia -> tool",
         '"cuanto es el monto de infima cuantia"',
         "Monto exacto 2025 ($7,263) via tool obtener_montos_pie.",
         "Monto correcto $7,263. Tool invocada. RAG no consultado.",
         ("PASA", VERDE)],

        ["CU-04",
         "Monto menor cuantia -> tool",
         '"hasta cuanto es menor cuantia de bienes"',
         "Monto $72,630 con base legal Art. 51 LOSNCP.",
         "Monto correcto $72,630. Art. 51 citado.",
         ("PASA", VERDE)],

        ["CU-05",
         "Plazo de impugnacion -> tool",
         '"cuantos dias tengo para impugnar"',
         "3 dias habiles desde notificacion. Art. 102 LOSNCP.",
         "Plazo correcto 3 dias habiles. Art. 102 LOSNCP citado.",
         ("PASA", VERDE)],

        ["CU-06",
         "Plazo firma de contrato -> tool",
         '"en cuantos dias se firma el contrato"',
         "15 dias habiles desde adjudicacion. Art. 69 LOSNCP.",
         "Plazo correcto 15 dias. Art. 69 citado.",
         ("PASA", VERDE)],

        ["CU-07",
         "Comparacion de procesos -> tools paralelas",
         '"diferencia entre cotizacion y menor cuantia"',
         "Tabla comparativa con montos de ambos procesos en USD.",
         "Comparacion correcta: $72,630 vs $72,630-$544,725. 3 tools invocadas.",
         ("PASA", VERDE)],

        ["CU-08",
         "Que proceso usar -> recomendar_tipo_contratacion",
         '"que proceso uso para comprar computadoras por $50,000"',
         "Recomendacion: Menor Cuantia bienes/servicios.",
         "Proceso correcto identificado. Monto y normativa correctos.",
         ("PASA", VERDE)],

        ["CU-09",
         "Como ser proveedor -> info_rup + RAG",
         '"quiero vender al estado, por donde empiezo"',
         "Pasos del RUP como primer paso obligatorio. Art. 16-18 LOSNCP.",
         "info_rup() invocada primero. Pasos correctos. Costo $0 citado.",
         ("PASA", VERDE)],

        ["CU-10",
         "Fecha actual -> tool fecha",
         '"que dia es hoy"',
         "Fecha y hora actual en Ecuador (UTC-5).",
         "Fecha correcta. Tool obtener_fecha_hora_ecuador invocada.",
         ("PASA", VERDE)],

        ["CU-11",
         "Mensaje multimedia (audio)",
         "Nota de voz de WhatsApp",
         "Respuesta orientativa sin invocar LLM. < 1s.",
         "Mensaje 'Solo puedo responder texto' enviado en < 800ms.",
         ("PASA", VERDE)],

        ["CU-12",
         "Mensaje duplicado (reintento Meta)",
         "Mismo mensaje_id enviado 2 veces",
         "Solo una respuesta entregada al usuario.",
         "Segunda solicitud ignorada. Log: 'Mensaje duplicado ignorado'.",
         ("PASA", VERDE)],

        ["CU-13",
         "Pregunta fuera de scope",
         '"cual es la capital de Francia"',
         "Redireccionar al scope de contratacion publica.",
         "Bot indica que solo atiende temas de contratacion publica.",
         ("PASA", VERDE)],

        ["CU-14",
         "Sigla SIE reconocida",
         '"que es la SIE"',
         "Respuesta sobre Subasta Inversa Electronica. Art. 47 LOSNCP.",
         "Sigla expandida correctamente. Art. 47 LOSNCP citado.",
         ("PASA", VERDE)],

        ["CU-15",
         "Envio de respuesta via Meta API",
         "Cualquier mensaje procesado correctamente",
         "Respuesta entregada al usuario por WhatsApp.",
         "PENDIENTE — graph.facebook.com bloqueado por Citrix DPI. Bot genera respuesta pero no puede enviarla.",
         ("PENDIENTE", NARANJA)],
    ],
    col_widths=[1.2, 3.5, 3.0, 3.8, 3.8, 1.2]
)

doc.add_paragraph()
subheading("5.2 Resumen de resultados")

add_table(
    ["Metrica", "Valor"],
    [
        ["Total casos evaluados",          "15"],
        ["Casos PASA",                     "14 (93.3%)"],
        ["Casos PENDIENTE",                "1 (6.7%) — entrega final por Citrix DPI"],
        ["Casos FALLA",                    "0"],
        ["Tiempo promedio de respuesta",   "~6 segundos (RAG + gemma4:e2b) / < 500ms (shortcuts)"],
        ["Precision de citacion legal",    "100% de casos con RAG activo"],
        ["Alucinacion de articulos",       "0 casos detectados"],
        ["Herramientas (tools) invocadas", "5 de 5 tools validadas"],
    ],
    col_widths=[6.5, 10.0]
)

doc.add_paragraph()
subheading("5.3 Pendiente critico")

body(
    "CU-15 — Entrega de respuesta via Meta API: el sistema genera y procesa la respuesta "
    "correctamente pero la entrega final al usuario de WhatsApp falla porque el firewall "
    "perimetral aplica inspeccion SSL (Citrix DPI) sobre el dominio graph.facebook.com, "
    "restableciendo la conexion TCP al momento del handshake TLS.",
    color=ROJO
)
body(
    "Accion requerida: Franklin Arias debe agregar graph.facebook.com a la lista de dominios "
    "excluidos de inspeccion SSL en el firewall Citrix. Una vez aplicado este cambio, "
    "el CU-15 quedara resuelto y el sistema estara 100% operativo.",
    bold=True
)

doc.add_paragraph()

# ── Pie ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "SERCOP — Coordinacion de Tecnologia de la Informacion | "
    "Anexos IT-DIO-2026-206 v1.0 | 16 abril 2026 | Uso interno"
)
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────
output = "docs/ANEXOS_IT-DIO-2026-206_SercoBot.docx"
doc.save(output)
print(f"OK Generado: {output}")
