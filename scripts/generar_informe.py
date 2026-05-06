"""
Genera el informe técnico IT-DIO-2026-206 en formato Word (.docx)
Uso: python scripts/generar_informe.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colores SERCOP ────────────────────────────────────────────
AZUL_SERCOP   = RGBColor(0x00, 0x33, 0x66)   # azul oscuro institucional
VERDE_SERCOP  = RGBColor(0x00, 0x70, 0x40)   # verde institucional
GRIS_TABLA    = RGBColor(0xD6, 0xE4, 0xF0)   # fondo cabecera tabla
NEGRO         = RGBColor(0x00, 0x00, 0x00)

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(table):
    """Aplica bordes simples a todas las celdas de la tabla."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)

def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = color or AZUL_SERCOP
    return p

def body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Cabecera
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        set_cell_bg(cell, "003366")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if r % 2 == 0:
                set_cell_bg(cell, "EAF2FA")

    set_cell_border(table)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ══════════════════════════════════════════════════════════════
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Fuente por defecto
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ── Encabezado institucional ──────────────────────────────────
header_table = doc.add_table(rows=2, cols=2)
header_table.cell(0, 0).merge(header_table.cell(1, 0))
header_table.cell(0, 0).text = "REPÚBLICA DEL ECUADOR\nSERCOP"
header_table.cell(0, 0).paragraphs[0].runs[0].bold = True
header_table.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = AZUL_SERCOP
header_table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)

header_table.cell(0, 1).text = "SISTEMA DE GESTIÓN"
header_table.cell(1, 1).text = "INFORME"
for c in [0, 1]:
    for r in [0, 1]:
        header_table.cell(r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ── Tabla de metadatos ────────────────────────────────────────
meta_table = doc.add_table(rows=3, cols=2)
meta_table.style = "Table Grid"
datos_meta = [
    ("SISTEMA DE GESTIÓN", "Vigencia: 2026/01/05"),
    ("INFORME",            "Versión: 22"),
    ("Informe Técnico Desarrollo del Chatbot", "Código: 7.5.P01.F02"),
]
for r, (izq, der) in enumerate(datos_meta):
    meta_table.cell(r, 0).text = izq
    meta_table.cell(r, 1).text = der
    meta_table.cell(r, 0).paragraphs[0].runs[0].font.size = Pt(9)
    meta_table.cell(r, 1).paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# Número y fecha
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r1 = p.add_run("NO. ")
r1.bold = True; r1.font.size = Pt(10)
r2 = p.add_run("IT-DIO-2026-206")
r2.font.size = Pt(10)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r3 = p2.add_run("FECHA: ")
r3.bold = True; r3.font.size = Pt(10)
r4 = p2.add_run("2026-04-16")
r4.font.size = Pt(10)

doc.add_paragraph()

# ── Título portada ────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = t.add_run("INFORME TÉCNICO DESARROLLO DEL CHATBOT\nDESARROLLO E INFRAESTRUCTURA — CASOS DE USO")
rt.bold = True
rt.font.size = Pt(16)
rt.font.color.rgb = AZUL_SERCOP

doc.add_paragraph()
doc.add_page_break()

# ── Tabla de contenido manual ─────────────────────────────────
heading(doc, "Contenido", level=1)
toc = [
    "1. ANTECEDENTES", "2. BASE LEGAL", "3. INTRODUCCIÓN",
    "4. DESARROLLO", "5. CONCLUSIONES Y RECOMENDACIONES",
    "6. ANEXOS", "7. FIRMAS DE ELABORACIÓN Y APROBACIÓN",
]
for item in toc:
    body(doc, item, space_after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. ANTECEDENTES
# ══════════════════════════════════════════════════════════════
heading(doc, "1.  ANTECEDENTES.")
body(doc, (
    "La Coordinación de Tecnología de la Información y Comunicaciones (TIC) del Servicio "
    "Nacional de Contratación Pública — SERCOP, identificó la necesidad de implementar un "
    "canal automatizado de atención ciudadana que permita resolver consultas frecuentes sobre "
    "normativa de contratación pública de forma ágil, precisa y disponible las 24 horas."
))
body(doc, (
    "Los ciudadanos, proveedores del Estado y funcionarios de entidades contratantes generan "
    "una alta demanda de consultas sobre la LOSNCP, su Reglamento General, resoluciones "
    "vigentes y procedimientos del portal ComprasPúblicas (SOCE), que supera la capacidad de "
    "atención directa de los canales institucionales en horario regular."
))
body(doc, (
    "En respuesta a esta necesidad, la Coordinación TIC desarrolló SercoBot, un asistente "
    "virtual de WhatsApp con inteligencia artificial, capaz de responder consultas normativas "
    "citando los artículos fuente, orientar a proveedores en su proceso de registro y "
    "participación, e informar sobre montos, plazos y tipos de procesos de contratación vigentes."
))
body(doc, (
    "El presente informe documenta el estado de desarrollo, la infraestructura implementada, "
    "los casos de uso validados y las incidencias identificadas durante las pruebas en "
    "producción realizadas el 15 de abril de 2026."
))

# ══════════════════════════════════════════════════════════════
# 2. BASE LEGAL
# ══════════════════════════════════════════════════════════════
heading(doc, "2.  BASE LEGAL.")
leyes = [
    "Ley Orgánica del Sistema Nacional de Contratación Pública — LOSNCP, Registro Oficial Suplemento N.° 140, 07 de octubre de 2025.",
    "Reglamento General a la LOSNCP — RGLOSNCP, versión octubre 2025.",
    "Estatuto Orgánico de Gestión Organizacional por Procesos del SERCOP, Resolución DSERCOP0001-2023.",
    "Norma Técnica Ecuatoriana NTE INEN-ISO 9001, Quinta edición, 2016-05 — Sistemas de Gestión de Calidad — Requisitos (ISO 9001:2015, IDT).",
    "Ley Orgánica de Telecomunicaciones, Registro Oficial Suplemento N.° 439, 18 de febrero de 2015.",
    "Ley Orgánica de Protección de Datos Personales, Registro Oficial Suplemento N.° 459, 26 de mayo de 2021.",
]
for ley in leyes:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(ley)
    run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 3. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════
heading(doc, "3.  INTRODUCCIÓN.")
body(doc, (
    "SercoBot es un asistente virtual institucional del SERCOP que opera sobre la plataforma "
    "de mensajería WhatsApp Business. Está diseñado para atender a ciudadanos, proveedores del "
    "Estado y funcionarios de entidades contratantes, respondiendo consultas sobre normativa de "
    "contratación pública con citación de artículos fuente, de forma automatizada, precisa y "
    "en lenguaje accesible."
))
body(doc, (
    "El sistema integra dos capas tecnológicas: (1) una capa de comunicación con WhatsApp "
    "basada en Meta Cloud API, y (2) una capa de inteligencia artificial que combina búsqueda "
    "semántica sobre base de datos vectorial (RAG) con herramientas estructuradas (tool calling) "
    "para datos exactos sobre montos, plazos y procedimientos."
))
body(doc, (
    "Principio rector: SercoBot nunca inventa artículos, montos ni plazos. Toda respuesta está "
    "respaldada por normativa indexada o herramientas con valores verificados. Cuando no dispone "
    "de información suficiente, orienta al usuario hacia los canales oficiales del SERCOP."
))

# ══════════════════════════════════════════════════════════════
# 4. DESARROLLO
# ══════════════════════════════════════════════════════════════
heading(doc, "4.  DESARROLLO.")

# 4.1 Arquitectura
heading(doc, "4.1  Arquitectura del Sistema", level=2)
body(doc, "Capa 1 — Comunicación WhatsApp (Meta Cloud API)", bold=True)
add_table(doc,
    ["Componente", "Tecnología"],
    [
        ["Servidor de aplicación", "FastAPI + Uvicorn (Python 3.9)"],
        ["Proveedor de mensajería", "Meta WhatsApp Cloud API v21.0"],
        ["Gestión de historial", "SQLAlchemy + PostgreSQL"],
        ["Despliegue", "Servidor Linux RHEL — infraestructura interna SERCOP"],
        ["Exposición HTTPS", "Citrix / NAT — sercobot.sercop.gob.ec"],
    ],
    col_widths=[6, 10]
)

body(doc, "Capa 2 — Inteligencia Artificial (RAG + Tool Calling)", bold=True)
add_table(doc,
    ["Componente", "Tecnología", "Estado"],
    [
        ["Modelo de lenguaje (LLM)", "Ollama — qwen2.5:3b (local)", "✅ Activo"],
        ["Modelo de embeddings", "nomic-embed-text 768d (Ollama)", "✅ Activo"],
        ["Base de datos vectorial", "PostgreSQL 16 + pgvector 0.8.2", "✅ Activo"],
        ["Búsqueda híbrida", "HNSW coseno + GIN tsvector español", "✅ Activo"],
        ["Fusión de rankings", "Reciprocal Rank Fusion (RRF) top-12", "✅ Activo"],
        ["Reranking semántico", "cross-encoder mmarco-mMiniLMv2 — top-4", "✅ Activo"],
    ],
    col_widths=[5, 8, 3]
)

body(doc, "Capa 3 — Infraestructura de Red y Datos (habilitada 15/04/2026)", bold=True)
add_table(doc,
    ["Componente", "Detalle"],
    [
        ["Dominio público", "sercobot.sercop.gob.ec"],
        ["IP pública", "157.100.62.125"],
        ["IP interna DMZ (red 100)", "192.168.100.131 / máscara 255.255.255.0 / GW 192.168.100.35"],
        ["Servidor de aplicación", "192.168.9.230 — puerto 8000"],
        ["Publicación", "Citrix — NAT firewall perimetral (HTTPS 443 + HTTP 80)"],
        ["Autorización", "Ing. Hugo Yépez — 2026-04-15 17:34"],
        ["Base de datos", "PostgreSQL 16 + pgvector 0.8.2 — 192.168.9.230 / sercop_db"],
    ],
    col_widths=[5.5, 10.5]
)

body(doc, "Flujo de red:", bold=True)
code_block(doc,
    "Usuario WhatsApp\n"
    "      ↓\n"
    "Meta Cloud API → sercobot.sercop.gob.ec\n"
    "      ↓ DNS → 157.100.62.125 (IP pública)\n"
    "      ↓ NAT firewall perimetral\n"
    "      192.168.100.131 (DMZ — Citrix)\n"
    "      ↓ reverse proxy\n"
    "      192.168.9.230:8000 (FastAPI — app-bdd-chatbot)"
)

# 4.2 Firewall
heading(doc, "4.2  Reglas de Firewall Habilitadas", level=2)
body(doc, "Salida desde servidor 192.168.9.230 — habilitadas por Franklin Arias el 15/04/2026:")
add_table(doc,
    ["Servicio", "Host de destino", "Puerto", "Protocolo"],
    [
        ["Groq API", "api.groq.com", "443", "HTTPS/TCP"],
        ["Meta WhatsApp", "graph.facebook.com", "443", "HTTPS/TCP"],
        ["Cloudflare Tunnel", "*.cloudflare.com", "443", "HTTPS/TCP"],
        ["GitHub", "github.com", "443", "HTTPS/TCP"],
        ["GitHub objetos", "objects.githubusercontent.com", "443", "HTTPS/TCP"],
        ["PyPI", "pypi.org", "443", "HTTPS/TCP"],
        ["PyPI archivos", "files.pythonhosted.org", "443", "HTTPS/TCP"],
        ["Ollama", "ollama.com", "443", "HTTPS/TCP"],
        ["HuggingFace", "huggingface.co", "443", "HTTPS/TCP"],
        ["HuggingFace CDN", "cdn-lfs.huggingface.co", "443", "HTTPS/TCP"],
    ],
    col_widths=[4, 6.5, 2, 3.5]
)

# 4.3 Base de conocimiento
heading(doc, "4.3  Base de Conocimiento Normativa", level=2)
body(doc, (
    "Se han indexado 3.059 fragmentos de 17 documentos oficiales del SERCOP en la base de "
    "datos vectorial sercop_db. El 91,5% de los fragmentos cuenta con metadata de artículo "
    "explícita, lo que permite citar la fuente legal con precisión en cada respuesta."
))
add_table(doc,
    ["Documento", "Tipo", "Chunks"],
    [
        ["Reglamento General LOSNCP — octubre 2025", "Reglamento", "1.132"],
        ["Normativa Secundaria de Contratación Pública 2025", "Reglamento", "1.125"],
        ["LOSNCP — RO 140, 07-X-2025", "Ley", "113"],
        ["Manual SOCE — Subasta Inversa Electrónica", "Manual", "108"],
        ["RE-SERCOP-2026-0001 (Metodología de Control)", "Resolución", "60"],
        ["RE-SERCOP-2026-0002 (Comité Interinstitucional)", "Resolución", "42"],
        ["Modelo de Pliego SICAE", "Resolución", "141"],
        ["Metodología de Control Final Sumillada", "Resolución", "137"],
        ["Código de Ética institucional", "Resolución", "89"],
        ["Norma Interna SICAE", "Resolución", "33"],
        ["Instructivo Extorsión", "Resolución", "29"],
        ["Resolución Régimen de Transición — nov. 2025", "Resolución", "14"],
        ["RE-SERCOP-2024-0144 / 2025-0152 / Fe de Errata", "Resoluciones", "11"],
        ["Manual SOCE — Fase Contractual Bienes y Servicios", "Manual", "19"],
        ["Glosario de Términos SERCOP", "Manual", "6"],
        ["TOTAL", "", "3.059"],
    ],
    col_widths=[9, 3, 2]
)

# 4.4 Tools
heading(doc, "4.4  Herramientas Estructuradas (Tool Calling)", level=2)
add_table(doc,
    ["Herramienta", "Activación", "Dato que entrega"],
    [
        ["obtener_montos_pie", "Pregunta por montos o umbrales en USD", "Umbrales exactos 2025/2026 por tipo de proceso"],
        ["recomendar_tipo_contratacion", "Pregunta sobre qué proceso usar", "Proceso recomendado con normativa aplicable"],
        ["obtener_plazos", "Pregunta sobre días, plazos, contratos, impugnaciones", "Plazos de 7 categorías"],
        ["info_rup", "Pregunta sobre Registro Único de Proveedores", "Requisitos, costo, tiempo, renovación"],
        ["obtener_fecha_hora_ecuador", "Pregunta sobre fecha u hora actual", "Fecha y hora Ecuador (UTC-5)"],
    ],
    col_widths=[4.5, 5.5, 6]
)

# 4.5 Casos de uso
heading(doc, "4.5  Casos de Uso Validados", level=2)
casos = [
    ("Consulta de monto para ínfima cuantía",
     "El sistema responde con el valor exacto del PIE 2026, cita Art. 52.1 LOSNCP y Art. 195 RGLOSNCP. Tiempo de respuesta: ~1,5 segundos."),
    ("Comparación de procesos de contratación",
     "Activa obtener_montos_pie y recomendar_tipo_contratacion, presenta tabla comparativa con umbrales en USD, plazos y nivel de competencia. Cita Art. 50 y 51 LOSNCP."),
    ("Plazo para impugnar una adjudicación",
     "Activa obtener_plazos('impugnacion'), responde con plazos exactos y base legal de LOSNCP y RGLOSNCP."),
    ("Registro como proveedor del Estado",
     "Activa info_rup() como primer paso, explica requisitos del RUP, cita Art. 16-18 LOSNCP."),
    ("Consulta normativa específica (garantías en licitación)",
     "Recupera chunks del RGLOSNCP vía RAG, responde con tipos de garantía, porcentajes y plazos con artículos específicos."),
]
for i, (titulo, desc) in enumerate(casos, 1):
    body(doc, f"Caso {i} — {titulo}", bold=True, space_after=2)
    body(doc, desc, space_after=8)

# 4.6 Mejoras implementadas
heading(doc, "4.6  Mejoras Implementadas en Producción (15/04/2026)", level=2)

body(doc, "a) Cambio de proveedor LLM — Groq → Ollama local", bold=True)
body(doc, (
    "Se configuró LLM_PROVIDER=ollama en el servidor de producción, utilizando el modelo "
    "qwen2.5:3b local. Los datos de consultas ciudadanas no salen de la infraestructura "
    "institucional del SERCOP."
))
add_table(doc,
    ["Parámetro", "Valor"],
    [
        ["LLM_PROVIDER", "ollama"],
        ["OLLAMA_MODEL", "qwen2.5:3b"],
        ["OLLAMA_NUM_CTX", "8192"],
        ["OLLAMA_MAX_TOKENS", "250"],
    ],
    col_widths=[6, 10]
)

body(doc, "b) Manejo de mensajes multimedia", bold=True)
body(doc, (
    "Se agregó respuesta automática para mensajes de tipo audio, imagen y sticker. "
    "El bot informa al usuario que solo procesa texto y sugiere ejemplos de consulta, "
    "sin pasar por el pipeline RAG+LLM."
))

body(doc, "c) Robustez en manejo de errores", bold=True)
body(doc, (
    "Se implementó registro detallado de errores con traceback completo y envío de "
    "mensaje informativo al usuario cuando el procesamiento falla, evitando silencio "
    "ante fallos técnicos."
))

# 4.7 Estado y pendientes
heading(doc, "4.7  Estado de Ítems — actualizado al 16/04/2026", level=2)
add_table(doc,
    ["Ítem", "Estado"],
    [
        ["Reglas firewall internet", "✅ Habilitado (Franklin Arias — 15/04/2026)"],
        ["DNS sercobot.sercop.gob.ec", "✅ Registrado (Rolando Coello — 15/04/2026)"],
        ["NAT / publicación Citrix", "✅ Configurado (15/04/2026)"],
        ["Webhook Meta Cloud API", "⚠️ En diagnóstico — error al enviar respuestas"],
        ["Git pull desde servidor", "⚠️ Bloqueado — Citrix inspecciona TLS hacia github.com"],
        ["Documentos faltantes (5 manuales SOCE + COA)", "🔧 En curso"],
        ["Evaluación RAGAS (faithfulness > 0,85)", "⏳ Pendiente"],
        ["Deploy gemma4:26b en servidor con GPU", "⏳ Pendiente"],
    ],
    col_widths=[9, 7]
)

body(doc, "Problema activo — envío de respuestas vía Meta API:", bold=True)
body(doc, (
    "Durante las pruebas del 15/04/2026 se registraron errores al procesar mensajes de "
    "usuarios reales. El LLM genera la respuesta correctamente pero se produce un fallo "
    "posterior. Se presume que Citrix aplica inspección SSL hacia graph.facebook.com, "
    "impidiendo que el servidor envíe las respuestas a WhatsApp. Diagnóstico en curso."
))

body(doc, "Restricción git en servidor:", bold=True)
body(doc, (
    "El servidor no puede ejecutar git pull desde GitHub debido a que la inspección SSL "
    "de Citrix interrumpe el handshake TLS. Se solicitará a Franklin Arias agregar "
    "github.com al bypass de inspección SSL. Mientras tanto, los despliegues se realizan "
    "mediante transferencia directa de archivos (scp)."
))

# ══════════════════════════════════════════════════════════════
# 5. CONCLUSIONES Y RECOMENDACIONES
# ══════════════════════════════════════════════════════════════
heading(doc, "5.  CONCLUSIONES Y RECOMENDACIONES.")
heading(doc, "Conclusiones:", level=2)
conclusiones = [
    ("1.", "El sistema SercoBot ha sido desarrollado e instalado satisfactoriamente en el servidor de producción del SERCOP (app-bdd-chatbot), con todas las funcionalidades core operativas: pipeline RAG, tool calling e integración con Meta WhatsApp Cloud API."),
    ("2.", "La infraestructura de red quedó habilitada el 15/04/2026 con la coordinación de los equipos de Redes, DNS, Firewall y Citrix del SERCOP, completando la publicación de sercobot.sercop.gob.ec con IP pública 157.100.62.125 y NAT institucional."),
    ("3.", "El sistema opera con LLM local (qwen2.5:3b vía Ollama), garantizando que los datos de consultas ciudadanas permanecen íntegramente en la infraestructura del SERCOP sin dependencia de servicios externos para la generación de respuestas."),
    ("4.", "Se identificó un problema en el envío de respuestas hacia Meta WhatsApp API, posiblemente relacionado con la inspección SSL de Citrix hacia graph.facebook.com. Este es el único componente pendiente para la operación plena del sistema."),
    ("5.", "La base de conocimiento cubre los documentos normativos de mayor demanda (LOSNCP, RGLOSNCP, resoluciones vigentes), con 3.059 fragmentos indexados y el 91,5% con metadata de artículo explícita."),
]
for num, texto in conclusiones:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(texto)
    run.font.size = Pt(10)

heading(doc, "Recomendaciones:", level=2)
recomendaciones = [
    "Inmediato: Solicitar a Franklin Arias agregar graph.facebook.com y github.com al bypass de inspección SSL del firewall perimetral. Sin esto, el bot no puede enviar respuestas a WhatsApp y el equipo de desarrollo no puede desplegar actualizaciones de código.",
    "Corto plazo: Una vez resuelto el envío a Meta, realizar prueba funcional completa con el número WhatsApp Business institucional.",
    "Corto plazo: Completar la base de conocimiento con los 5 manuales SOCE pendientes y la resolución oficial de montos PIE 2026.",
    "Mediano plazo: Ejecutar evaluación de calidad del sistema RAG mediante la suite RAGAS con mínimo 20 preguntas representativas — meta: faithfulness > 0,85.",
    "Mediano plazo: Evaluar upgrade del modelo LLM local a gemma4:e2b o superior cuando se confirme disponibilidad de RAM en el servidor de producción.",
]
for rec in recomendaciones:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(rec)
    run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 6. ANEXOS
# ══════════════════════════════════════════════════════════════
heading(doc, "6.  ANEXOS.")
anexos = [
    "Anexo 1: Arquitectura técnica detallada del sistema SercoBot (diagrama de componentes)",
    "Anexo 2: Listado completo de documentos normativos indexados con número de fragmentos",
    "Anexo 3: Informe de diagnóstico de conectividad — reporte_firewall.txt (13/04/2026)",
    "Anexo 4: Ticket MANTIS de habilitación de infraestructura (15/04/2026)",
    "Anexo 5: Resultados de pruebas funcionales — casos de uso validados",
]
for a in anexos:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(a)
    run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 7. FIRMAS
# ══════════════════════════════════════════════════════════════
heading(doc, "7.  FIRMAS DE ELABORACIÓN Y APROBACIÓN")
firmas_table = doc.add_table(rows=4, cols=4)
firmas_table.style = "Table Grid"

headers_firmas = ["Nombre / Cargo", "Rol", "Fecha", "Firma de Aceptación"]
for i, h in enumerate(headers_firmas):
    cell = firmas_table.cell(0, i)
    cell.text = h
    set_cell_bg(cell, "003366")
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

firmas = [
    ("Mauricio Ruiz\nAnalista de Desarrollo TIC — SERCOP", "Elaborado", "2026-04-16", ""),
    ("Ing. Paúl Vásquez Méndez\nDirector de Infraestructura y Operaciones TIC — SERCOP", "Revisado", "", ""),
    ("", "Aprobado", "", ""),
]
for r, (nombre, rol, fecha, firma) in enumerate(firmas, 1):
    firmas_table.cell(r, 0).text = nombre
    firmas_table.cell(r, 1).text = rol
    firmas_table.cell(r, 2).text = fecha
    firmas_table.cell(r, 3).text = firma
    for c in range(4):
        cell = firmas_table.cell(r, c)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.height = Cm(2)

set_cell_border(firmas_table)

# Pie de página
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Código: 7.5.P01.F02 — Versión 22 — Vigencia: 2026/01/05\n"
    "Dirección: Plataforma Gubernamental Financiera, Amazonas entre Unión Nacional de Periodistas "
    "y Alfonso Pereira, Bloque Amarillo, Piso 7 — Quito, Ecuador"
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Guardar ───────────────────────────────────────────────────
output = "docs/IT-DIO-2026-206_SercoBot.docx"
doc.save(output)
print(f"OK Informe generado: {output}")
